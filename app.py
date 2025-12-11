# app.py
import io
import json
import re

import joblib
import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
import PyPDF2
import plotly.express as px

# Page configuration
st.set_page_config(
    page_title="Document Classification System",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown(
    """
    <style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


@st.cache_resource
def load_models():
    """Load trained models and encoders"""
    try:
        model = joblib.load("best_model.pkl")
        vectorizer = joblib.load("tfidf_vectorizer.pkl")
        label_encoder = joblib.load("label_encoder.pkl")

        with open("model_info.json", "r", encoding="utf-8") as f:
            model_info = json.load(f)

        return model, vectorizer, label_encoder, model_info
    except Exception as e:
        st.error(f"Error loading models: {e}")
        return None, None, None, None


def preprocess_text(text: str) -> str:
    """Clean and preprocess text"""
    if not isinstance(text, str):
        return ""
    text = text.lower()
    text = re.sub(r"[^a-zA-Z\s]", " ", text)
    text = " ".join(text.split())
    return text


def extract_text_from_docx(file) -> str:
    """
    Extract text from Word document.
    Accepts a Streamlit uploaded file (UploadedFile) or a local file-like object.
    """
    try:
        # Try reading directly (python-docx accepts file-like objects)
        file.seek(0)
        doc = Document(file)
        text = "\n".join([p.text for p in doc.paragraphs])
        if text.strip():
            return text
    except Exception:
        # Fallback: read bytes and reopen via BytesIO
        try:
            file.seek(0)
            raw = file.read()
            doc = Document(io.BytesIO(raw))
            text = "\n".join([p.text for p in doc.paragraphs])
            return text
        except Exception as exc:
            st.warning(f"⚠️ Could not read Word document: {getattr(file, 'name', 'uploaded_file')}. {exc}")
            return ""
    return ""


def extract_text_from_pdf(file) -> str:
    """
    Extract text from PDF file using PyPDF2.
    Accepts a Streamlit uploaded file (UploadedFile) or a local file-like object.
    """
    try:
        file.seek(0)
        # PyPDF2 accepts a file-like object
        reader = PyPDF2.PdfReader(file)
        text_parts = []
        for page in reader.pages:
            # extract_text may return None for some pages, guard it
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        # Some environments need bytes read and BytesIO
        try:
            file.seek(0)
            raw = file.read()
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as exc:
            st.error(f"Error reading PDF: {exc}")
            return ""


def predict_document(text: str, model, vectorizer, label_encoder):
    """Predict document category and return category + dict of probs"""
    try:
        cleaned_text = preprocess_text(text)
        if not cleaned_text:
            return None, {}

        # Vectorize
        text_vectorized = vectorizer.transform([cleaned_text])

        # Predict label index (or encoded label)
        prediction = model.predict(text_vectorized)[0]

        # Get probabilities for each class
        probabilities = None
        if hasattr(model, "predict_proba"):
            probabilities = model.predict_proba(text_vectorized)[0]
        elif hasattr(model, "decision_function"):
            # fallback: convert decision_function to pseudo probs (not ideal)
            raw_scores = model.decision_function(text_vectorized)[0]
            # softmax
            exp = np.exp(raw_scores - np.max(raw_scores))
            probabilities = exp / exp.sum()
        else:
            # unknown model type: set uniform distribution
            n_classes = len(label_encoder.classes_)
            probabilities = np.ones(n_classes) / n_classes

        # Convert class indices to labels safely
        class_indices = np.arange(len(probabilities))
        class_labels = label_encoder.inverse_transform(class_indices)

        # Build probabilities dict
        prob_dict = {lbl: float(prob) for lbl, prob in zip(class_labels, probabilities)}

        # Convert predicted label back to original label if prediction is encoded index or label
        try:
            # if prediction is an index (int), inverse_transform will map it
            predicted_label = label_encoder.inverse_transform([prediction])[0]
        except Exception:
            # fallback: if prediction already a string label
            predicted_label = prediction

        return predicted_label, prob_dict

    except Exception as e:
        st.error(f"Prediction error: {e}")
        return None, {}


def main():
    # Header
    st.markdown('<h1 class="main-header">📄 Document Classification System</h1>', unsafe_allow_html=True)
    st.markdown("### Automated Resume Classification for HRM")

    # Load models
    model, vectorizer, label_encoder, model_info = load_models()

    if model is None or vectorizer is None or label_encoder is None:
        st.error("⚠️ Models not found or failed to load. Please ensure the files exist in this folder:")
        st.info("Required files: best_model.pkl, tfidf_vectorizer.pkl, label_encoder.pkl, model_info.json")
        return

    # Sidebar: model info (safe access)
    with st.sidebar:
        st.header("📊 Model Information")
        try:
            st.metric("Best Model", model_info.get("best_model", "N/A"))
            st.metric("Accuracy", f"{model_info.get('accuracy', 0):.2%}")
            st.metric("F1-Score", f"{model_info.get('f1_score', 0):.2%}")
        except Exception:
            st.write("Model info not available or malformed.")

        st.markdown("---")
        st.header("📁 Categories")
        try:
            for i, cat in enumerate(model_info.get("categories", list(label_encoder.classes_)), 1):
                st.write(f"{i}. {cat}")
        except Exception:
            for i, cat in enumerate(label_encoder.classes_, 1):
                st.write(f"{i}. {cat}")

        st.markdown("---")
        st.header("ℹ️ About")
        st.info(
            "This system automatically classifies resumes into different "
            "categories to reduce manual effort in HRM processes."
        )

    # Main content tabs
    tab1, tab2, tab3 = st.tabs(["📤 Upload Document", "📝 Text Input", "📈 Batch Processing"])

    # Tab 1: File Upload
    with tab1:
        st.header("Upload Resume Document")
        uploaded_file = st.file_uploader("Choose a file (PDF or Word)", type=["pdf", "docx", "doc"], help="Upload a resume document for classification")

        if uploaded_file is not None:
            # Extract text
            filename = uploaded_file.name.lower()
            if filename.endswith(".pdf"):
                text = extract_text_from_pdf(uploaded_file)
            else:
                text = extract_text_from_docx(uploaded_file)

            if text and text.strip():
                st.success(f"✅ File loaded: {uploaded_file.name}")

                # Show preview
                with st.expander("📄 Document Preview"):
                    st.text_area("Content", text[:1000] + "..." if len(text) > 1000 else text, height=200)

                # Predict button
                if st.button("🔍 Classify Document"):
                    with st.spinner("Analyzing document..."):
                        category, probabilities = predict_document(text, model, vectorizer, label_encoder)

                    if category:
                        st.markdown("---")
                        st.header("🎯 Classification Results")
                        col1, col2 = st.columns([1, 2])

                        with col1:
                            st.markdown("### Predicted Category")
                            st.markdown(f"<h2 style='color: #1f77b4;'>{category}</h2>", unsafe_allow_html=True)
                            conf = probabilities.get(category, 0.0)
                            st.metric("Confidence", f"{conf:.2%}")

                        with col2:
                            st.markdown("### Probability Distribution")
                            prob_df = pd.DataFrame({"Category": list(probabilities.keys()), "Probability": list(probabilities.values())})
                            prob_df = prob_df.sort_values("Probability", ascending=True)
                            fig = px.bar(prob_df, x="Probability", y="Category", orientation="h", color="Probability")
                            fig.update_layout(height=300, showlegend=False)
                            st.plotly_chart(fig, use_container_width=True)
                    else:
                        st.error("Could not predict the document. Make sure the document contains readable text.")
            else:
                st.error("❌ Could not extract text from the document. Try a different file or make sure the file contains selectable text.")

    # Tab 2: Text Input
    with tab2:
        st.header("Enter Resume Text")
        text_input = st.text_area("Paste resume content here", height=300, placeholder="Enter or paste resume text here...")

        if st.button("🔍 Classify Text"):
            if text_input and text_input.strip():
                with st.spinner("Analyzing text..."):
                    category, probabilities = predict_document(text_input, model, vectorizer, label_encoder)

                if category:
                    st.markdown("---")
                    st.header("🎯 Classification Results")
                    col1, col2 = st.columns(2)

                    with col1:
                        st.markdown("### Predicted Category")
                        st.markdown(f"<h2 style='color: #1f77b4;'>{category}</h2>", unsafe_allow_html=True)
                        st.metric("Confidence", f"{probabilities.get(category, 0.0):.2%}")

                    with col2:
                        st.markdown("### All Probabilities")
                        prob_df = pd.DataFrame({"Category": list(probabilities.keys()), "Probability": [f'{p:.2%}' for p in probabilities.values()]})
                        prob_df = prob_df.sort_values("Probability", ascending=False)
                        st.dataframe(prob_df, hide_index=True, use_container_width=True)
                else:
                    st.error("Prediction failed. Check input or model compatibility.")
            else:
                st.warning("⚠️ Please enter some text")

    # Tab 3: Batch Processing
    with tab3:
        st.header("Batch Document Processing")
        st.info("Upload multiple documents for batch classification")
        uploaded_files = st.file_uploader("Choose multiple files", type=["pdf", "docx", "doc"], accept_multiple_files=True)

        if uploaded_files:
            st.write(f"📁 {len(uploaded_files)} files uploaded")

            if st.button("🚀 Process All Documents"):
                results = []
                progress_bar = st.progress(0)
                status_text = st.empty()

                total = len(uploaded_files)
                for idx, file in enumerate(uploaded_files, start=1):
                    status_text.text(f"Processing {file.name} ({idx}/{total})...")

                    # Extract text
                    name = file.name.lower()
                    if name.endswith(".pdf"):
                        text = extract_text_from_pdf(file)
                    else:
                        text = extract_text_from_docx(file)

                    if text and text.strip():
                        category, probabilities = predict_document(text, model, vectorizer, label_encoder)
                        conf = probabilities.get(category, 0.0) if category else 0.0
                        results.append({"Filename": file.name, "Predicted Category": category or "N/A", "Confidence": f"{conf:.2%}"})
                    else:
                        results.append({"Filename": file.name, "Predicted Category": "N/A", "Confidence": "0.00%"})

                    # update progress as percent
                    progress_bar.progress(int((idx / total) * 100))

                status_text.text("✅ Processing complete!")

                # Display results
                st.markdown("---")
                st.header("📊 Batch Results")
                results_df = pd.DataFrame(results)
                st.dataframe(results_df, hide_index=True, use_container_width=True)

                # Summary statistics
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Total Processed", len(results_df))
                with col2:
                    try:
                        most_common = results_df.loc[results_df["Predicted Category"] != "N/A", "Predicted Category"].mode()[0]
                    except Exception:
                        most_common = "N/A"
                    st.metric("Most Common Category", most_common)
                with col3:
                    try:
                        avg_conf = results_df["Confidence"].str.rstrip("%").astype(float).mean()
                        st.metric("Avg Confidence", f"{avg_conf:.1f}%")
                    except Exception:
                        st.metric("Avg Confidence", "0%")

                # Category distribution
                st.markdown("### Category Distribution")
                try:
                    cat_counts = results_df.loc[results_df["Predicted Category"] != "N/A", "Predicted Category"].value_counts()
                    fig = px.pie(values=cat_counts.values, names=cat_counts.index, title="Distribution of Predicted Categories")
                    st.plotly_chart(fig, use_container_width=True)
                except Exception:
                    st.info("No valid categories to show distribution.")

                # Download results
                csv = results_df.to_csv(index=False)
                st.download_button(label="📥 Download Results as CSV", data=csv, file_name="classification_results.csv", mime="text/csv")


if __name__ == "__main__":
    main()
